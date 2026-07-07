"""Extract DOCX paragraphs, tables, images, and embedded files."""

from __future__ import annotations

import argparse
import io
import re
import shutil
import zipfile
from pathlib import Path

from docx import Document

OLE_HEADER = bytes.fromhex("d0cf11e0a1b11ae1")
USEFUL_ASCII_TERMS = {"APP", "WEB", "H5", "API", "SQL", "HTTP", "HTTPS", "JSON", "XML"}


def clean(text: str) -> str:
    return " ".join((text or "").split())


def is_useful_term(term: str) -> bool:
    return bool(re.search(r"[\u4e00-\u9fff]", term)) or term.upper() in USEFUL_ASCII_TERMS


def readable_chinese_terms(data: bytes) -> list[str]:
    terms: list[str] = []
    for encoding in ("utf-16le", "latin1"):
        text = data.decode(encoding, errors="ignore")
        for term in re.findall(r"[\u4e00-\u9fffA-Za-z0-9_\-（）()/]{2,}", text):
            term = term.strip()
            if term and is_useful_term(term) and term not in terms:
                terms.append(term)
    return terms


def try_read_ole_package(data: bytes) -> bytes | None:
    try:
        import olefile  # type: ignore
    except Exception:
        return None

    if not data.startswith(OLE_HEADER):
        return None

    try:
        ole = olefile.OleFileIO(io.BytesIO(data))
        if ole.exists("Package"):
            return ole.openstream("Package").read()
    except Exception:
        return None

    return None


def analyze_embedding(asset: Path) -> list[str]:
    lines: list[str] = []
    data = asset.read_bytes()

    if not data.startswith(OLE_HEADER):
        return lines

    lines.append(f"[OLE] {asset} type=oleObject")
    package = try_read_ole_package(data)
    if package is None:
        lines.append(f"[OLE] {asset} package=not_extracted reason=olefile_missing_or_unreadable")
        return lines

    if zipfile.is_zipfile(io.BytesIO(package)):
        package_path = asset.with_suffix(".vsdx")
        package_path.write_bytes(package)
        lines.append(f"[OLE_PACKAGE] {asset} extracted={package_path} package=zip")
        if package_path.suffix.lower() == ".vsdx":
            lines.extend(analyze_vsdx(package_path))
    else:
        package_path = asset.with_suffix(".package")
        package_path.write_bytes(package)
        lines.append(f"[OLE_PACKAGE] {asset} extracted={package_path} package=binary")

    return lines


def analyze_vsdx(vsdx_path: Path) -> list[str]:
    lines: list[str] = [f"[VSDX] {vsdx_path}"]
    terms: list[str] = []

    with zipfile.ZipFile(vsdx_path) as archive:
        names = archive.namelist()
        lines.append(f"[VSDX] entries={len(names)}")
        for name in names:
            if name.endswith(".xml"):
                text = archive.read(name).decode("utf-8", errors="ignore")
                plain = re.sub(r"<[^>]+>", " ", text)
                for term in re.findall(r"[\u4e00-\u9fffA-Za-z0-9_\-（）()/]{2,}", plain):
                    term = term.strip()
                    if term and is_useful_term(term) and term not in terms:
                        terms.append(term)
            elif name.lower().endswith((".emf", ".png", ".jpg", ".jpeg")):
                data = archive.read(name)
                for term in readable_chinese_terms(data):
                    if term not in terms:
                        terms.append(term)

    if terms:
        lines.append("[VSDX_TERMS] " + "、".join(terms[:120]))
    else:
        lines.append("[VSDX_TERMS] none")

    return lines


def export_assets(input_path: Path, asset_dir: Path) -> list[Path]:
    exported: list[Path] = []
    prefixes = ("word/media/", "word/embeddings/")

    with zipfile.ZipFile(input_path) as archive:
        for member in archive.namelist():
            if not member.startswith(prefixes) or member.endswith("/"):
                continue
            relative = Path(member)
            target = asset_dir / relative.parent.name / relative.name
            target.parent.mkdir(parents=True, exist_ok=True)
            with archive.open(member) as source, target.open("wb") as destination:
                shutil.copyfileobj(source, destination)
            exported.append(target)

    return exported


def extract_docx(input_path: Path, asset_dir: Path | None = None, analyze_embeddings: bool = True) -> list[str]:
    doc = Document(input_path)
    lines: list[str] = [f"source={input_path}", f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}"]

    if asset_dir is not None:
        assets = export_assets(input_path, asset_dir)
        lines.append(f"assets={asset_dir} count={len(assets)}")
        for asset in assets:
            lines.append(f"[ASSET] {asset}")
            if analyze_embeddings and asset.suffix.lower() == ".bin":
                lines.extend(analyze_embedding(asset))

    for index, paragraph in enumerate(doc.paragraphs):
        text = clean(paragraph.text)
        if not text:
            continue
        style = paragraph.style.name if paragraph.style else ""
        lines.append(f"[P{index:04d}][{style}] {text}")

    for table_index, table in enumerate(doc.tables):
        lines.append("")
        lines.append(f"[TABLE {table_index}] rows={len(table.rows)} cols={len(table.columns)}")
        for row_index, row in enumerate(table.rows):
            cells = [clean(cell.text) for cell in row.cells]
            if any(cells):
                lines.append(f"[T{table_index}R{row_index}] " + " | ".join(cells))

    return lines


def main() -> None:
    parser = argparse.ArgumentParser(description="Extract a DOCX file into UTF-8 text.")
    parser.add_argument("input", type=Path)
    parser.add_argument("--out", type=Path, required=True)
    parser.add_argument("--asset-dir", type=Path, help="Directory for extracted word/media and word/embeddings files.")
    parser.add_argument("--no-analyze-embeddings", action="store_true", help="Skip OLE .bin package extraction and text hint analysis.")
    args = parser.parse_args()

    asset_dir = args.asset_dir
    if asset_dir is None:
        asset_dir = args.out.parent / f"{args.input.stem}-assets"

    lines = extract_docx(args.input, asset_dir, analyze_embeddings=not args.no_analyze_embeddings)
    args.out.parent.mkdir(parents=True, exist_ok=True)
    args.out.write_text("\n".join(lines), encoding="utf-8")
    print(args.out)


if __name__ == "__main__":
    main()
